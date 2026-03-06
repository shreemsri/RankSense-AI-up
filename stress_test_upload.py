import asyncio
import aiohttp
import os
from docx import Document
import time

async def upload_resume(session, filename, user_id):
    url = "http://localhost:8000/process"
    data = aiohttp.FormData()
    data.add_field('file', open(filename, 'rb'), filename=filename)
    data.add_field('user_id', user_id)
    
    print(f"Starting upload for {filename}...")
    start_time = time.time()
    async with session.post(url, data=data) as response:
        status = response.status
        content = await response.json()
        elapsed = time.time() - start_time
        print(f"Finished {filename}: Status {status}, Time {elapsed:.2f}s")
        return status, content

async def main():
    # 1. Create a dummy docx
    filename = "stress_test_resume.docx"
    doc = Document()
    doc.add_heading('John Doe Resume', 0)
    doc.add_paragraph('Software Engineer with 5 years of experience in Python and FastAPI.')
    doc.add_paragraph('GitHub: github.com/johndoe')
    doc.save(filename)
    
    # 2. Upload concurrently
    async with aiohttp.ClientSession() as session:
        tasks = []
        for i in range(5):
            tasks.append(upload_resume(session, f"stress_{i}.docx", "stress_test_user"))
            # Create unique files to avoid filename collision (though INSERT OR REPLACE should handle it)
            # Actually, let's use the SAME filename to test the UNIQUE constraint fix too.
            # No, let's test BOTH scenarios.
        
        # Scenario 1: Same filename, same user (Concurrent INSERT OR REPLACE)
        print("\n--- SCENARIO 1: SAME FILENAME CONCURRENT ---")
        tasks = [upload_resume(session, filename, "stress_test_user") for _ in range(3)]
        results = await asyncio.gather(*tasks)
        
        # Scenario 2: Different filenames, same user (Concurrent AI processing)
        print("\n--- SCENARIO 2: DIFFERENT FILENAMES CONCURRENT ---")
        for i in range(5):
            fname = f"stress_{i}.docx"
            os.rename(filename, fname) if i == 0 else doc.save(fname)
        
        tasks = [upload_resume(session, f"stress_{i}.docx", "stress_test_user") for i in range(5)]
        results = await asyncio.gather(*tasks)

    # Cleanup
    for i in range(5):
        if os.path.exists(f"stress_{i}.docx"): os.remove(f"stress_{i}.docx")
    if os.path.exists(filename): os.remove(filename)

if __name__ == "__main__":
    asyncio.run(main())
